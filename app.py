import streamlit as st
import cv2
import numpy as np
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import time

# Configure page with custom styling
st.set_page_config(
    page_title="Transcribe — Image to Word",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for a distinctive, refined interface
st.markdown("""
<style>
    /* Import distinctive fonts */
    @import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@300;400;600&family=JetBrains+Mono:wght@300;400&family=Karla:wght@400;600&display=swap');
    
    /* Hide Streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Reset and base styles */
    .stApp {
        background: linear-gradient(165deg, #fafaf9 0%, #f5f5f4 50%, #e7e5e4 100%);
        font-family: 'Karla', sans-serif;
    }
    
    /* Main container - centered, refined layout */
    .main .block-container {
        max-width: 1100px;
        padding: 3rem 2rem;
    }
    
    /* Custom header */
    .custom-header {
        text-align: center;
        margin-bottom: 4rem;
        animation: fadeInDown 0.8s ease-out;
    }
    
    .custom-header h1 {
        font-family: 'Cormorant Garamond', serif;
        font-size: 3.5rem;
        font-weight: 300;
        letter-spacing: -0.02em;
        color: #1c1917;
        margin: 0;
        line-height: 1.1;
    }
    
    .custom-header .subtitle {
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.75rem;
        letter-spacing: 0.1em;
        text-transform: uppercase;
        color: #78716c;
        margin-top: 0.5rem;
        font-weight: 300;
    }
    
    /* Stage-based layout */
    .stage-container {
        display: flex;
        gap: 1px;
        margin: 3rem 0;
        background: #d6d3d1;
        border-radius: 2px;
        overflow: hidden;
    }
    
    .stage {
        flex: 1;
        background: white;
        padding: 2rem;
        position: relative;
        transition: all 0.3s ease;
        min-height: 400px;
    }
    
    .stage.active {
        background: #fafaf9;
        flex: 1.5;
    }
    
    .stage-number {
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.65rem;
        color: #a8a29e;
        letter-spacing: 0.1em;
        margin-bottom: 1rem;
        font-weight: 300;
    }
    
    .stage-title {
        font-family: 'Cormorant Garamond', serif;
        font-size: 1.5rem;
        font-weight: 400;
        color: #1c1917;
        margin-bottom: 1.5rem;
        letter-spacing: -0.01em;
    }
    
    /* Upload zone */
    .upload-zone {
        border: 1px solid #e7e5e4;
        border-radius: 2px;
        padding: 3rem 2rem;
        text-align: center;
        background: white;
        cursor: pointer;
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .upload-zone:hover {
        border-color: #78716c;
        background: #fafaf9;
    }
    
    .upload-zone::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(120, 113, 108, 0.05), transparent);
        transition: left 0.5s ease;
    }
    
    .upload-zone:hover::before {
        left: 100%;
    }
    
    /* Image preview */
    .image-preview {
        border: 1px solid #e7e5e4;
        border-radius: 2px;
        overflow: hidden;
        margin: 1rem 0;
        animation: fadeIn 0.5s ease-out;
    }
    
    /* Processing state */
    .processing-indicator {
        text-align: center;
        padding: 3rem 2rem;
    }
    
    .processing-dots {
        display: flex;
        gap: 0.5rem;
        justify-content: center;
        margin: 2rem 0;
    }
    
    .processing-dots .dot {
        width: 8px;
        height: 8px;
        border-radius: 50%;
        background: #78716c;
        animation: pulse 1.5s ease-in-out infinite;
    }
    
    .processing-dots .dot:nth-child(2) {
        animation-delay: 0.2s;
    }
    
    .processing-dots .dot:nth-child(3) {
        animation-delay: 0.4s;
    }
    
    /* Results display */
    .results-container {
        animation: fadeIn 0.6s ease-out;
    }
    
    .metric-card {
        background: white;
        border: 1px solid #e7e5e4;
        border-radius: 2px;
        padding: 1.5rem;
        margin: 1rem 0;
    }
    
    .metric-value {
        font-family: 'Cormorant Garamond', serif;
        font-size: 2.5rem;
        font-weight: 300;
        color: #1c1917;
        line-height: 1;
    }
    
    .metric-label {
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.65rem;
        color: #78716c;
        letter-spacing: 0.1em;
        text-transform: uppercase;
        margin-top: 0.5rem;
    }
    
    /* Custom button */
    .stButton > button {
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.75rem;
        letter-spacing: 0.05em;
        text-transform: uppercase;
        background: #1c1917;
        color: white;
        border: none;
        border-radius: 2px;
        padding: 1rem 2rem;
        cursor: pointer;
        transition: all 0.3s ease;
        width: 100%;
        font-weight: 400;
    }
    
    .stButton > button:hover {
        background: #292524;
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(28, 25, 23, 0.15);
    }
    
    /* Download button special styling */
    .stDownloadButton > button {
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.75rem;
        letter-spacing: 0.05em;
        text-transform: uppercase;
        background: white;
        color: #1c1917;
        border: 1px solid #1c1917;
        border-radius: 2px;
        padding: 1rem 2rem;
        cursor: pointer;
        transition: all 0.3s ease;
        width: 100%;
        font-weight: 400;
    }
    
    .stDownloadButton > button:hover {
        background: #1c1917;
        color: white;
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(28, 25, 23, 0.15);
    }
    
    /* Text preview */
    .text-preview {
        font-family: 'Karla', sans-serif;
        font-size: 0.9rem;
        line-height: 1.7;
        color: #44403c;
        background: white;
        border: 1px solid #e7e5e4;
        border-radius: 2px;
        padding: 2rem;
        margin: 1rem 0;
        max-height: 300px;
        overflow-y: auto;
    }
    
    /* Progress indicator */
    .custom-progress {
        height: 2px;
        background: #e7e5e4;
        border-radius: 2px;
        overflow: hidden;
        margin: 2rem 0;
    }
    
    .custom-progress-bar {
        height: 100%;
        background: #1c1917;
        transition: width 0.3s ease;
    }
    
    /* Info cards */
    .info-card {
        background: white;
        border-left: 2px solid #78716c;
        padding: 1.5rem;
        margin: 1rem 0;
        font-size: 0.9rem;
        line-height: 1.6;
        color: #57534e;
    }
    
    /* Animations */
    @keyframes fadeIn {
        from {
            opacity: 0;
            transform: translateY(10px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    @keyframes fadeInDown {
        from {
            opacity: 0;
            transform: translateY(-20px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    @keyframes pulse {
        0%, 100% {
            opacity: 0.3;
            transform: scale(0.8);
        }
        50% {
            opacity: 1;
            transform: scale(1);
        }
    }
    
    /* Expander styling */
    .streamlit-expanderHeader {
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.75rem;
        letter-spacing: 0.05em;
        text-transform: uppercase;
        color: #78716c;
        border: 1px solid #e7e5e4;
        border-radius: 2px;
        background: white;
    }
    
    .streamlit-expanderContent {
        border: 1px solid #e7e5e4;
        border-top: none;
        border-radius: 0 0 2px 2px;
        background: #fafaf9;
    }
    
    /* File uploader */
    .stFileUploader {
        border: 1px solid #e7e5e4;
        border-radius: 2px;
        padding: 2rem;
        background: white;
    }
    
    /* Success/Info messages */
    .stSuccess, .stInfo {
        background: white;
        border-left: 2px solid #1c1917;
        padding: 1rem 1.5rem;
        border-radius: 2px;
        font-family: 'Karla', sans-serif;
        font-size: 0.9rem;
    }
    
    /* Footer */
    .custom-footer {
        margin-top: 6rem;
        padding-top: 3rem;
        border-top: 1px solid #e7e5e4;
        text-align: center;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.65rem;
        color: #a8a29e;
        letter-spacing: 0.1em;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'stage' not in st.session_state:
    st.session_state.stage = 'upload'  # upload, processing, results
if 'processed_data' not in st.session_state:
    st.session_state.processed_data = None

def preprocess_image(image):
    """Preprocess the image for better OCR results"""
    img_array = np.array(image)
    if len(img_array.shape) == 3:
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
    else:
        gray = img_array
    _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    denoised = cv2.fastNlMeansDenoising(binary, None, 10, 7, 21)
    kernel = np.ones((1, 1), np.uint8)
    processed = cv2.dilate(denoised, kernel, iterations=1)
    processed = cv2.erode(processed, kernel, iterations=1)
    return processed

def detect_formatting(text, confidence):
    """Detect basic formatting from OCR results"""
    formatting = {
        'is_bold': False,
        'is_italic': False,
        'is_heading': False,
        'alignment': 'left'
    }
    if text.isupper() and len(text.split()) <= 10:
        formatting['is_heading'] = True
        formatting['is_bold'] = True
    heading_patterns = [r'^Chapter \d+', r'^Section \d+', r'^\d+\.', r'^[A-Z][A-Z\s]+$']
    for pattern in heading_patterns:
        if re.match(pattern, text.strip()):
            formatting['is_heading'] = True
            formatting['is_bold'] = True
            break
    if len(text.strip()) < 50 and len(text.strip().split()) <= 8:
        words = text.strip().split()
        if len(words) > 0 and words[0][0].isupper():
            formatting['alignment'] = 'center'
    return formatting

def extract_text_with_formatting(image):
    """Extract text from image using Tesseract with formatting information"""
    processed_img = preprocess_image(image)
    pil_img = Image.fromarray(processed_img)
    ocr_data = pytesseract.image_to_data(pil_img, output_type=pytesseract.Output.DICT)
    lines = []
    current_line = {'text': '', 'confidence': [], 'block_num': -1, 'line_num': -1}
    n_boxes = len(ocr_data['text'])
    for i in range(n_boxes):
        text = ocr_data['text'][i].strip()
        conf = int(ocr_data['conf'][i])
        block_num = ocr_data['block_num'][i]
        line_num = ocr_data['line_num'][i]
        if conf > 0 and text:
            if current_line['line_num'] != line_num or current_line['block_num'] != block_num:
                if current_line['text']:
                    lines.append(current_line)
                current_line = {'text': text, 'confidence': [conf], 'block_num': block_num, 'line_num': line_num}
            else:
                current_line['text'] += ' ' + text
                current_line['confidence'].append(conf)
    if current_line['text']:
        lines.append(current_line)
    return lines

def create_word_document(lines_data):
    """Create a Word document with formatting preserved"""
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    prev_block = -1
    for line_data in lines_data:
        text = line_data['text'].strip()
        if not text:
            continue
        avg_confidence = sum(line_data['confidence']) / len(line_data['confidence']) if line_data['confidence'] else 0
        formatting = detect_formatting(text, avg_confidence)
        if prev_block != -1 and line_data['block_num'] != prev_block:
            doc.add_paragraph()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        if formatting['is_bold']:
            run.bold = True
        if formatting['is_italic']:
            run.italic = True
        if formatting['is_heading']:
            run.font.size = Pt(16)
        else:
            run.font.size = Pt(11)
        if formatting['alignment'] == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif formatting['alignment'] == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        prev_block = line_data['block_num']
    return doc

# Header
st.markdown("""
<div class="custom-header">
    <h1>Transcribe</h1>
    <div class="subtitle">Image to Word Conversion</div>
</div>
""", unsafe_allow_html=True)

# Main interface based on stage
uploaded_file = st.file_uploader(
    "Select Document Image",
    type=['png', 'jpg', 'jpeg', 'bmp', 'tiff'],
    label_visibility="collapsed"
)

# Stage 1: Upload
if uploaded_file is not None and st.session_state.stage == 'upload':
    st.markdown('<div class="stage-number">STAGE 01 — INPUT</div>', unsafe_allow_html=True)
    st.markdown('<div class="stage-title">Document Preview</div>', unsafe_allow_html=True)
    
    image = Image.open(uploaded_file)
    st.image(image, use_container_width=True)
    
    st.markdown(f"""
    <div class="info-card">
        Image dimensions: {image.size[0]} × {image.size[1]} pixels
    </div>
    """, unsafe_allow_html=True)
    
    if st.button("Begin Transcription", key="start_process"):
        st.session_state.stage = 'processing'
        st.rerun()

# Stage 2: Processing
elif uploaded_file is not None and st.session_state.stage == 'processing':
    st.markdown('<div class="stage-number">STAGE 02 — PROCESSING</div>', unsafe_allow_html=True)
    st.markdown('<div class="stage-title">Extracting Text</div>', unsafe_allow_html=True)
    
    # Processing animation
    st.markdown("""
    <div class="processing-indicator">
        <div class="processing-dots">
            <div class="dot"></div>
            <div class="dot"></div>
            <div class="dot"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Actual processing
    image = Image.open(uploaded_file)
    
    progress_container = st.empty()
    status_container = st.empty()
    
    status_container.markdown('<div class="metric-label">Preprocessing image...</div>', unsafe_allow_html=True)
    time.sleep(0.5)
    
    status_container.markdown('<div class="metric-label">Analyzing document structure...</div>', unsafe_allow_html=True)
    lines_data = extract_text_with_formatting(image)
    time.sleep(0.5)
    
    status_container.markdown('<div class="metric-label">Generating formatted document...</div>', unsafe_allow_html=True)
    doc = create_word_document(lines_data)
    time.sleep(0.5)
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    # Store in session state
    st.session_state.processed_data = {
        'doc_bytes': doc_bytes,
        'lines_data': lines_data,
        'image': image
    }
    
    st.session_state.stage = 'results'
    st.rerun()

# Stage 3: Results
elif uploaded_file is not None and st.session_state.stage == 'results' and st.session_state.processed_data:
    st.markdown('<div class="stage-number">STAGE 03 — OUTPUT</div>', unsafe_allow_html=True)
    st.markdown('<div class="stage-title">Transcription Complete</div>', unsafe_allow_html=True)
    
    lines_data = st.session_state.processed_data['lines_data']
    doc_bytes = st.session_state.processed_data['doc_bytes']
    
    # Statistics in elegant metric cards
    col1, col2, col3 = st.columns(3)
    
    total_words = sum([len(line['text'].split()) for line in lines_data])
    avg_confidence = sum([sum(line['confidence'])/len(line['confidence']) for line in lines_data if line['confidence']]) / len(lines_data) if lines_data else 0
    
    with col1:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{len(lines_data)}</div>
            <div class="metric-label">Lines Detected</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{total_words}</div>
            <div class="metric-label">Words Extracted</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{avg_confidence:.0f}%</div>
            <div class="metric-label">Confidence</div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Text preview
    with st.expander("Preview Extracted Text", expanded=False):
        preview_text = "\n\n".join([line['text'] for line in lines_data[:15]])
        st.markdown(f'<div class="text-preview">{preview_text}</div>', unsafe_allow_html=True)
        if len(lines_data) > 15:
            st.markdown(f'<div class="metric-label">...and {len(lines_data) - 15} more lines</div>', unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Download button
    st.download_button(
        label="Download Word Document",
        data=doc_bytes,
        file_name="transcribed_document.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Reset button
    if st.button("Transcribe Another Document", key="reset"):
        st.session_state.stage = 'upload'
        st.session_state.processed_data = None
        st.rerun()

# Initial state
elif uploaded_file is None:
    st.markdown("""
    <div class="info-card">
        Upload a clear image of a document to begin transcription. Supported formats: JPG, PNG, TIFF, BMP. 
        For best results, use high-resolution images with good contrast and lighting.
    </div>
    """, unsafe_allow_html=True)

# Footer
st.markdown("""
<div class="custom-footer">
    Powered by Tesseract OCR — Built with Precision
</div>
""", unsafe_allow_html=True)


# import streamlit as st
# import cv2
# import numpy as np
# from PIL import Image
# import pytesseract
# from docx import Document
# from docx.shared import Pt, RGBColor, Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# import io
# import re
# import tempfile
# import os

# # Configure page
# st.set_page_config(
#     page_title="Image to Word Converter",
#     page_icon="📄",
#     layout="wide"
# )

# # Title and description
# st.title("📄 Image to Word Converter")
# st.markdown("""
# Convert scanned documents or images to editable Word format while preserving basic formatting.
# Supports JPG, PNG, and other common image formats.
# """)

# def preprocess_image(image):
#     """
#     Preprocess the image for better OCR results
#     """
#     # Convert PIL Image to numpy array
#     img_array = np.array(image)
    
#     # Convert to grayscale
#     if len(img_array.shape) == 3:
#         gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
#     else:
#         gray = img_array
    
#     # Apply thresholding to get binary image
#     _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    
#     # Denoise
#     denoised = cv2.fastNlMeansDenoising(binary, None, 10, 7, 21)
    
#     # Dilation and erosion to remove noise
#     kernel = np.ones((1, 1), np.uint8)
#     processed = cv2.dilate(denoised, kernel, iterations=1)
#     processed = cv2.erode(processed, kernel, iterations=1)
    
#     return processed

# def detect_formatting(text, confidence):
#     """
#     Detect basic formatting from OCR results
#     Returns formatting attributes for text
#     """
#     formatting = {
#         'is_bold': False,
#         'is_italic': False,
#         'is_heading': False,
#         'alignment': 'left'
#     }
    
#     # Simple heuristics for formatting detection
#     # Check if text is all uppercase (likely heading)
#     if text.isupper() and len(text.split()) <= 10:
#         formatting['is_heading'] = True
#         formatting['is_bold'] = True
    
#     # Check if text starts with common heading patterns
#     heading_patterns = [r'^Chapter \d+', r'^Section \d+', r'^\d+\.', r'^[A-Z][A-Z\s]+$']
#     for pattern in heading_patterns:
#         if re.match(pattern, text.strip()):
#             formatting['is_heading'] = True
#             formatting['is_bold'] = True
#             break
    
#     # Check for centered text (short lines, common in titles)
#     if len(text.strip()) < 50 and len(text.strip().split()) <= 8:
#         words = text.strip().split()
#         if len(words) > 0 and words[0][0].isupper():
#             formatting['alignment'] = 'center'
    
#     return formatting

# def extract_text_with_formatting(image):
#     """
#     Extract text from image using Tesseract with formatting information
#     """
#     # Preprocess image
#     processed_img = preprocess_image(image)
    
#     # Convert to PIL Image for pytesseract
#     pil_img = Image.fromarray(processed_img)
    
#     # Get detailed OCR data
#     ocr_data = pytesseract.image_to_data(pil_img, output_type=pytesseract.Output.DICT)
    
#     # Group text by lines
#     lines = []
#     current_line = {
#         'text': '',
#         'confidence': [],
#         'block_num': -1,
#         'line_num': -1
#     }
    
#     n_boxes = len(ocr_data['text'])
#     for i in range(n_boxes):
#         text = ocr_data['text'][i].strip()
#         conf = int(ocr_data['conf'][i])
#         block_num = ocr_data['block_num'][i]
#         line_num = ocr_data['line_num'][i]
        
#         if conf > 0 and text:  # Only consider confident detections
#             # New line detected
#             if current_line['line_num'] != line_num or current_line['block_num'] != block_num:
#                 if current_line['text']:
#                     lines.append(current_line)
#                 current_line = {
#                     'text': text,
#                     'confidence': [conf],
#                     'block_num': block_num,
#                     'line_num': line_num
#                 }
#             else:
#                 current_line['text'] += ' ' + text
#                 current_line['confidence'].append(conf)
    
#     # Add last line
#     if current_line['text']:
#         lines.append(current_line)
    
#     return lines

# def create_word_document(lines_data):
#     """
#     Create a Word document with formatting preserved
#     """
#     doc = Document()
    
#     # Set document margins
#     sections = doc.sections
#     for section in sections:
#         section.top_margin = Inches(1)
#         section.bottom_margin = Inches(1)
#         section.left_margin = Inches(1)
#         section.right_margin = Inches(1)
    
#     prev_block = -1
    
#     for line_data in lines_data:
#         text = line_data['text'].strip()
#         if not text:
#             continue
        
#         avg_confidence = sum(line_data['confidence']) / len(line_data['confidence']) if line_data['confidence'] else 0
        
#         # Detect formatting
#         formatting = detect_formatting(text, avg_confidence)
        
#         # Add spacing between blocks
#         if prev_block != -1 and line_data['block_num'] != prev_block:
#             doc.add_paragraph()
        
#         # Add paragraph
#         paragraph = doc.add_paragraph()
#         run = paragraph.add_run(text)
        
#         # Apply formatting
#         if formatting['is_bold']:
#             run.bold = True
        
#         if formatting['is_italic']:
#             run.italic = True
        
#         # Set font size
#         if formatting['is_heading']:
#             run.font.size = Pt(16)
#         else:
#             run.font.size = Pt(11)
        
#         # Set alignment
#         if formatting['alignment'] == 'center':
#             paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         elif formatting['alignment'] == 'right':
#             paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#         else:
#             paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
#         prev_block = line_data['block_num']
    
#     return doc

# def main():
#     # Sidebar for information
#     with st.sidebar:
#         st.header("ℹ️ About")
#         st.markdown("""
#         **Features:**
#         - Image preprocessing for better OCR
#         - Text extraction with Tesseract OCR
#         - Automatic formatting detection
#         - Preserves bold, italic, headings
#         - Detects text alignment
#         - Generates formatted .docx files
        
#         **Supported Formats:**
#         - JPG, JPEG, PNG, BMP, TIFF
        
#         **Tips:**
#         - Use clear, high-resolution images
#         - Ensure good lighting and contrast
#         - Avoid skewed or rotated images
#         """)
        
#         st.header("🎯 Project Info")
#         st.markdown("""
#         **Task 1: Image-to-Word Converter (MVP)**
        
#         A desktop application that extracts text from images 
#         while preserving basic formatting.
#         """)
    
#     # Main content
#     col1, col2 = st.columns([1, 1])
    
#     with col1:
#         st.header("📤 Upload Image")
#         uploaded_file = st.file_uploader(
#             "Choose an image file",
#             type=['png', 'jpg', 'jpeg', 'bmp', 'tiff'],
#             help="Upload a clear image of a document"
#         )
        
#         if uploaded_file is not None:
#             # Display uploaded image
#             image = Image.open(uploaded_file)
#             st.image(image, caption="Uploaded Image", use_container_width=True)
            
#             # Show image info
#             st.info(f"📏 Image Size: {image.size[0]} x {image.size[1]} pixels")
    
#     with col2:
#         st.header("⚙️ Processing")
        
#         if uploaded_file is not None:
#             process_button = st.button("🚀 Convert to Word", type="primary", use_container_width=True)
            
#             if process_button:
#                 with st.spinner("Processing image... This may take a moment."):
#                     try:
#                         # Load image
#                         image = Image.open(uploaded_file)
                        
#                         # Progress tracking
#                         progress_bar = st.progress(0)
#                         status_text = st.empty()
                        
#                         # Step 1: Preprocess
#                         status_text.text("Step 1/3: Preprocessing image...")
#                         progress_bar.progress(33)
                        
#                         # Step 2: Extract text
#                         status_text.text("Step 2/3: Extracting text with OCR...")
#                         lines_data = extract_text_with_formatting(image)
#                         progress_bar.progress(66)
                        
#                         # Step 3: Create document
#                         status_text.text("Step 3/3: Creating Word document...")
#                         doc = create_word_document(lines_data)
#                         progress_bar.progress(100)
                        
#                         # Save to bytes
#                         doc_bytes = io.BytesIO()
#                         doc.save(doc_bytes)
#                         doc_bytes.seek(0)
                        
#                         status_text.text("✅ Conversion complete!")
                        
#                         # Success message
#                         st.success("🎉 Document created successfully!")
                        
#                         # Display extracted text preview
#                         with st.expander("📝 Preview Extracted Text"):
#                             preview_text = "\n\n".join([line['text'] for line in lines_data[:10]])
#                             st.text_area("First 10 lines:", preview_text, height=200)
#                             if len(lines_data) > 10:
#                                 st.info(f"... and {len(lines_data) - 10} more lines")
                        
#                         # Statistics
#                         total_words = sum([len(line['text'].split()) for line in lines_data])
#                         avg_confidence = sum([sum(line['confidence'])/len(line['confidence']) for line in lines_data if line['confidence']]) / len(lines_data) if lines_data else 0
                        
#                         col_a, col_b, col_c = st.columns(3)
#                         col_a.metric("Lines", len(lines_data))
#                         col_b.metric("Words", total_words)
#                         col_c.metric("Confidence", f"{avg_confidence:.1f}%")
                        
#                         # Download button
#                         st.download_button(
#                             label="📥 Download Word Document",
#                             data=doc_bytes,
#                             file_name="converted_document.docx",
#                             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                             use_container_width=True
#                         )
                        
#                     except Exception as e:
#                         st.error(f"❌ Error during conversion: {str(e)}")
#                         st.info("💡 Try uploading a clearer image or check if Tesseract OCR is properly installed.")
#         else:
#             st.info("👈 Please upload an image to begin conversion")
    
#     # Footer
#     st.markdown("---")
#     st.markdown("""
#     <div style='text-align: center'>
#         <p>Built with ❤️ using Streamlit, OpenCV, and Tesseract OCR</p>
#         <p><small>Image-to-Word Converter MVP v1.0</small></p>
#     </div>
#     """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
